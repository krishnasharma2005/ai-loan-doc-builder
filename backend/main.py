from fastapi import FastAPI
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from openai import OpenAI
import os

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React frontend
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

class LoanInput(BaseModel):
    borrower: str
    amount: float
    tenure: int
    interest_rate: float
    collateral: str
    jurisdiction: str

# ---------- AI Clause Generation ----------
def generate_ai_clause(loan_data):
    try:
        prompt = f"""
        Generate legally correct loan clauses for:
        Borrower: {loan_data.borrower}
        Amount: {loan_data.amount}
        Collateral: {loan_data.collateral}
        Jurisdiction: {loan_data.jurisdiction}

        Include:
        - repayment clause
        - penalty clause
        - default clause
        """

        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}]
        )

        return resp.choices[0].message["content"]

    except Exception as e:
        # 🔴 FALLBACK WHEN API FAILS
        return (
            "Repayment Clause: The borrower shall repay the loan in equal monthly "
            "installments over the agreed tenure.\n\n"
            "Penalty Clause: Any delay in repayment shall attract a penalty interest "
            "as per lender policy.\n\n"
            "Default Clause: Failure to repay installments for 90 days shall be "
            "considered an event of default."
        )


# ---------- Generate Document API ----------
@app.post("/generate-document")
def generate_document(data: LoanInput):

    # 1️⃣ Generate AI clauses
    ai_clauses = generate_ai_clause(data)

    # 2️⃣ Create document
    doc = Document()
    doc.add_heading("Loan Agreement", level=1)

    doc.add_paragraph(f"Borrower: {data.borrower}")
    doc.add_paragraph(f"Loan Amount: {data.amount}")
    doc.add_paragraph(f"Tenure: {data.tenure} months")
    doc.add_paragraph(f"Interest Rate: {data.interest_rate}%")
    doc.add_paragraph(f"Collateral: {data.collateral}")
    doc.add_paragraph(f"Jurisdiction: {data.jurisdiction}")

    doc.add_heading("AI-Generated Clauses", level=2)
    doc.add_paragraph(ai_clauses)

    # 3️⃣ Save document
    file_path = "loan_output.docx"
    doc.save(file_path)

    # 4️⃣ Return AI text (frontend preview uses this)
    return {
        "message": "Document created",
        "file": file_path,
        "ai_clauses": ai_clauses
    }


# ---------- Validation Model ----------
class DocumentText(BaseModel):
    text: str


# ---------- Validation API ----------
@app.post("/validate")
def validate_document(payload: DocumentText):
    text = payload.text

    issues = []

    # Very simple rule-based validator
    if "penalty" not in text.lower():
        issues.append({"issue": "Penalty clause missing", "severity": "High"})

    if "repayment" not in text.lower():
        issues.append({"issue": "Repayment clause missing", "severity": "High"})

    if len(text) < 200:
        issues.append({"issue": "Document appears too short", "severity": "Medium"})

    if "witness" not in text.lower():
        issues.append({"issue": "Witness signature line missing", "severity": "Low"})

    return {"issues": issues}
 
@app.get("/download")
def download_document():
    file_path = "loan_output.docx"

    if not os.path.exists(file_path):
        return {"error": "File not found"}

    return FileResponse(
        path=file_path,
        filename="Loan_Agreement.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )